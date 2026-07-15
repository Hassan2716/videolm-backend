"""
Export Service — Generates PDF, DOCX, TXT, SRT, JSON, CSV, ZIP exports.
All formats use free/open-source libraries only.
"""
import os, uuid, json, csv, zipfile
from datetime import datetime
from typing import List
from loguru import logger
from sqlalchemy.orm import Session

from app.core.database import SessionLocal, ExportJob, Project, Transcript, Summary, Frame, ChatMessage, StudyAsset
from app.core.config import settings


class ExportService:

    @staticmethod
    def generate(export_id: str, project_id: str, fmt: str, content_types: List[str]):
        db = SessionLocal()
        try:
            # Update status
            db.query(ExportJob).filter(ExportJob.id == export_id).update({"status": "processing"})
            db.commit()

            out_dir = os.path.join(settings.export_dir, project_id)
            os.makedirs(out_dir, exist_ok=True)

            project = db.query(Project).filter(Project.id == project_id).first()
            transcript = db.query(Transcript).filter(Transcript.project_id == project_id).first()
            summaries = db.query(Summary).filter(Summary.project_id == project_id).all()
            frames = db.query(Frame).filter(Frame.project_id == project_id).order_by(Frame.timestamp_seconds).all()
            chat_msgs = db.query(ChatMessage).filter(ChatMessage.project_id == project_id).order_by(ChatMessage.created_at).all()
            study_assets = db.query(StudyAsset).filter(StudyAsset.project_id == project_id).all()

            data = {
                "project": project, "transcript": transcript,
                "summaries": summaries, "frames": frames,
                "chat_msgs": chat_msgs, "study_assets": study_assets,
                "content_types": content_types,
            }

            dispatch = {
                "pdf": ExportService._pdf,
                "docx": ExportService._docx,
                "txt": ExportService._txt,
                "json": ExportService._json,
                "srt": ExportService._srt,
                "vtt": ExportService._vtt,
                "csv": ExportService._csv,
                "zip": ExportService._zip,
            }

            fn = dispatch.get(fmt)
            if not fn:
                raise ValueError(f"Unsupported format: {fmt}")

            file_path = fn(out_dir, project_id, data)

            db.query(ExportJob).filter(ExportJob.id == export_id).update({
                "status": "complete", "file_path": file_path,
            })
            db.commit()
            logger.info(f"Export complete: {file_path}")

        except Exception as e:
            logger.error(f"Export failed: {e}")
            db.query(ExportJob).filter(ExportJob.id == export_id).update({
                "status": "failed",
            })
            db.commit()
        finally:
            db.close()

    # ── PDF ──────────────────────────────────────────────────────────────

    @staticmethod
    def _pdf(out_dir: str, project_id: str, data: dict) -> str:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            HRFlowable, PageBreak, KeepTogether,
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

        path = os.path.join(out_dir, f"report_{project_id[:8]}.pdf")
        doc = SimpleDocTemplate(
            path, pagesize=A4,
            topMargin=20*mm, bottomMargin=20*mm,
            leftMargin=18*mm, rightMargin=18*mm
        )

        # ── Styles ──────────────────────────────────────────────────────
        styles = getSampleStyleSheet()
        BLUE    = colors.HexColor("#1e3a8a")
        LBLUE   = colors.HexColor("#2563eb")
        SLATE   = colors.HexColor("#64748b")
        BODY    = colors.HexColor("#1e293b")
        LIGHT   = colors.HexColor("#f8fafc")
        BORDER  = colors.HexColor("#e2e8f0")

        title_s = ParagraphStyle("TT", parent=styles["Title"],   fontSize=22, textColor=BLUE,   spaceAfter=6,  alignment=TA_CENTER)
        sub_s   = ParagraphStyle("SS", parent=styles["Normal"],  fontSize=10, textColor=SLATE,  spaceAfter=3,  alignment=TA_CENTER)
        h1_s    = ParagraphStyle("H1", parent=styles["Heading1"],fontSize=16, textColor=BLUE,   spaceBefore=14, spaceAfter=6)
        h2_s    = ParagraphStyle("H2", parent=styles["Heading2"],fontSize=13, textColor=LBLUE,  spaceBefore=10, spaceAfter=4)
        body_s  = ParagraphStyle("BB", parent=styles["Normal"],  fontSize=10, leading=15, textColor=BODY, spaceAfter=4, alignment=TA_JUSTIFY)
        mono_s  = ParagraphStyle("MM", parent=styles["Normal"],  fontSize=9,  leading=13, textColor=BODY, fontName="Courier")
        label_s = ParagraphStyle("LL", parent=styles["Normal"],  fontSize=9,  textColor=SLATE, leftIndent=8)
        ts_s    = ParagraphStyle("TS", parent=styles["Normal"],  fontSize=9,  leading=13, textColor=BODY)

        ct      = data["content_types"]
        project = data["project"]
        story   = []

        def hr():
            story.append(Spacer(1, 2*mm))
            story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
            story.append(Spacer(1, 2*mm))

        def section_header(title):
            story.append(PageBreak())
            story.append(Paragraph(title, h1_s))
            story.append(HRFlowable(width="100%", thickness=1.5, color=LBLUE))
            story.append(Spacer(1, 4*mm))

        def safe_text(txt):
            if not txt: return ""
            return str(txt).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

        # ── Cover page ───────────────────────────────────────────────────
        if "cover" in ct or not ct:
            story.append(Spacer(1, 30*mm))
            story.append(Paragraph("VideoLM", title_s))
            story.append(Paragraph("AI Video Analysis Report", ParagraphStyle("ST", parent=sub_s, fontSize=14, textColor=LBLUE, spaceAfter=12)))
            hr()
            src = project.title or project.source_url or project.source_filename or "Video Analysis"
            story.append(Paragraph(f"Source: {safe_text(src)}", sub_s))
            story.append(Paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", sub_s))
            if project.duration_seconds:
                m, s = int(project.duration_seconds // 60), int(project.duration_seconds % 60)
                story.append(Paragraph(f"Duration: {m}:{s:02d}", sub_s))
            story.append(Spacer(1, 6*mm))

        # ── Video Metadata ───────────────────────────────────────────────
        if "metadata" in ct:
            section_header("Video Metadata")
            meta_rows = [
                ["Field", "Value"],
                ["Source", safe_text(project.source_url or project.source_filename or "—")],
                ["Status", str(project.status or "—")],
                ["Duration", f"{int((project.duration_seconds or 0)//60)}:{int((project.duration_seconds or 0)%60):02d}" if project.duration_seconds else "—"],
                ["Language", str(getattr(data.get("transcript"), "language", "—") or "—")],
                ["Word Count", str(getattr(data.get("transcript"), "word_count", "—") or "—")],
                ["Frames Extracted", str(len(data["frames"]))],
                ["Created", str(project.created_at)[:19] if project.created_at else "—"],
            ]
            t = Table(meta_rows, colWidths=[55*mm, 115*mm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0),(-1,0), BLUE),
                ("TEXTCOLOR",  (0,0),(-1,0), colors.white),
                ("FONTSIZE",   (0,0),(-1,-1), 10),
                ("FONTNAME",   (0,0),(-1,0), "Helvetica-Bold"),
                ("GRID",       (0,0),(-1,-1), 0.5, BORDER),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white, LIGHT]),
                ("VALIGN",     (0,0),(-1,-1), "MIDDLE"),
                ("TOPPADDING", (0,0),(-1,-1), 5),
                ("BOTTOMPADDING",(0,0),(-1,-1), 5),
            ]))
            story.append(t)

        # ── Summaries ────────────────────────────────────────────────────
        TYPE_MAP = {
            "short_summary":    ("short",    "Short Summary (Executive Overview)"),
            "medium_summary":   ("medium",   "Study Notes"),
            "detailed_summary": ("detailed", "Detailed Notes"),
            "bullet_summary":   ("bullets",  "Key Points (Bullets)"),
            "academic_summary": ("academic", "Academic Summary"),
        }
        for ct_key, (stype, label) in TYPE_MAP.items():
            if ct_key in ct:
                matches = [s for s in data["summaries"] if s.summary_type == stype]
                if matches:
                    section_header(label)
                    for line in matches[0].content.split("\n"):
                        line = line.strip()
                        if not line: continue
                        if line.startswith("## "):
                            story.append(Paragraph(safe_text(line[3:]), h2_s))
                        elif line.startswith("**") and line.endswith("**"):
                            story.append(Paragraph(f"<b>{safe_text(line[2:-2])}</b>", body_s))
                        elif line.startswith("• "):
                            story.append(Paragraph(f"• {safe_text(line[2:])}", body_s))
                        else:
                            story.append(Paragraph(safe_text(line), body_s))

        # ── Full Transcript ───────────────────────────────────────────────
        transcript = data["transcript"]
        if "transcript" in ct and transcript:
            section_header("Full Transcript")
            for seg in (transcript.segments or []):
                ts = seg.get("start", 0)
                m2, s2 = int(ts // 60), int(ts % 60)
                story.append(Paragraph(
                    f'<font color="#2563eb">[{m2:02d}:{s2:02d}]</font> {safe_text(seg.get("text", ""))}',
                    ts_s
                ))

        # ── Frame Captions ───────────────────────────────────────────────
        if "frame_captions" in ct and data["frames"]:
            section_header("Frame Captions & OCR")
            for frame in data["frames"]:
                cap = frame.caption or ""
                ocr = frame.ocr_text or ""
                if not cap and not ocr: continue
                ts_label = frame.timestamp_label or "—"
                story.append(Paragraph(
                    f'<b><font color="#2563eb">[{ts_label}]</font></b> ' + safe_text(cap),
                    body_s
                ))
                if ocr:
                    story.append(Paragraph(
                        f"OCR: {safe_text(ocr[:300])}",
                        label_s
                    ))
                story.append(Spacer(1, 1.5*mm))

        # ── Key Timestamps ───────────────────────────────────────────────
        if "key_timestamps" in ct and data["frames"]:
            section_header("Key Timestamps")
            ts_rows = [["Time", "Visual Type", "Caption"]]
            for frame in data["frames"]:
                cap = (frame.caption or "")[:80]
                vtype = (frame.visual_type or "unknown").title()
                ts_rows.append([frame.timestamp_label or "—", vtype, safe_text(cap)])
                t = Table(ts_rows, colWidths=[22*mm, 35*mm, 113*mm], repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND",  (0,0),(-1,0),  BLUE),
                ("TEXTCOLOR",   (0,0),(-1,0),  colors.white),
                ("FONTNAME",    (0,0),(-1,0),  "Helvetica-Bold"),
                ("FONTSIZE",    (0,0),(-1,-1), 9),
                ("GRID",        (0,0),(-1,-1), 0.5, BORDER),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white, LIGHT]),
                ("VALIGN",      (0,0),(-1,-1), "TOP"),
                ("TOPPADDING",  (0,0),(-1,-1), 6),
                ("BOTTOMPADDING",(0,0),(-1,-1), 6),
                ("WORDWRAP",    (0,0),(-1,-1), True),
            ]))
            story.append(t)

        # ── Flashcards ───────────────────────────────────────────────────
        if "flashcards" in ct:
            flash_assets = [a for a in data["study_assets"] if a.asset_type == "flashcard"]
            if flash_assets:
                section_header("Flashcards")
                cards = flash_assets[0].content.get("flashcards", [])

                if cards:
                    flash_style = ParagraphStyle(
                        "FlashStyle",
                        parent=body_s,
                        fontName="Helvetica",
                        fontSize=8,
                        leading=11,
                        alignment=TA_LEFT,
                        wordWrap="CJK",      # forces a break even on long unbroken tokens/URLs
                        spaceBefore=0,
                        spaceAfter=0,
                    )
                    flash_header_style = ParagraphStyle(
                        "FlashHeaderStyle",
                        parent=flash_style,
                        textColor=colors.white,
                        fontName="Helvetica-Bold",
                    )

                    fc_rows = [[
                        Paragraph("Concept / Question", flash_header_style),
                        Paragraph("Answer / Definition", flash_header_style),
                    ]]

                    for card in cards:
                        front = str(
                            card.get("front")
                            or card.get("question")
                            or card.get("concept", "")
                        ).strip()
                        back = str(
                            card.get("back")
                            or card.get("answer", "")
                        ).strip()

                        if front or back:
                            fc_rows.append([
                                Paragraph(safe_text(front), flash_style),
                                Paragraph(safe_text(back), flash_style),
                            ])

                    # Columns must add up to the same usable width as your other
                    # tables (page width minus margins). Using Paragraph cells
                    # (not raw strings) is what makes long answers wrap instead
                    # of overflowing past the column/table border.
                    t = Table(
                        fc_rows,
                        colWidths=[55 * mm, 115 * mm],
                        repeatRows=1,
                        splitByRow=True,
                    )

                    t.setStyle(TableStyle([
                        ("BACKGROUND",    (0,0),(-1,0), BLUE),
                        ("TEXTCOLOR",     (0,0),(-1,0), colors.white),
                        ("FONTNAME",      (0,0),(-1,0), "Helvetica-Bold"),
                        ("FONTSIZE",      (0,0),(-1,-1), 8),
                        ("GRID",          (0,0),(-1,-1), 0.5, BORDER),
                        ("ROWBACKGROUNDS",(0,1),(-1,-1), [colors.white, LIGHT]),
                        ("VALIGN",        (0,0),(-1,-1), "TOP"),
                        ("ALIGN",         (0,0),(-1,-1), "LEFT"),
                        ("TOPPADDING",    (0,0),(-1,-1), 6),
                        ("BOTTOMPADDING", (0,0),(-1,-1), 6),
                        ("LEFTPADDING",   (0,0),(-1,-1), 5),
                        ("RIGHTPADDING",  (0,0),(-1,-1), 5),
                    ]))

                    story.append(t)
                else:
                    story.append(
                        Paragraph(
                            "No flashcards generated yet. Go to the Flashcards tab and click Generate.",
                            body_s,
                        )
                    )

        # ── Quiz Questions ───────────────────────────────────────────────
        if "quiz" in ct:
            quiz_assets = [a for a in data["study_assets"] if a.asset_type == "quiz"]
            if quiz_assets:
                section_header("Quiz Questions")
                questions = quiz_assets[0].content.get("questions", [])
                if questions:
                    for i, q in enumerate(questions, 1):
                        qtext = safe_text(q.get("question",""))
                        diff  = q.get("difficulty","medium").title()
                        qtype = q.get("type","mcq").replace("_"," ").title()
                        story.append(Paragraph(
                            f"<b>Q{i}. [{diff} · {qtype}]</b> {qtext}", body_s
                        ))
                        # Options for MCQ
                        opts = q.get("options", {})
                        if opts:
                            for letter, opt_text in opts.items():
                                is_correct = letter == q.get("correct","")
                                marker = "✓ " if is_correct else "   "
                                story.append(Paragraph(
                                    f"{marker}<b>{letter})</b> {safe_text(str(opt_text)[:150])}",
                                    ParagraphStyle("OPT", parent=body_s, leftIndent=12, fontSize=9,
                                                   textColor=colors.HexColor("#16a34a") if is_correct else BODY)
                                ))
                        # True/False — FIXED
                        elif q.get("answer") is not None:
                            answer_text = safe_text(str(q.get('answer', '')))
                            story.append(Paragraph(
                                f"   Answer: <b>{answer_text}</b>",
                                ParagraphStyle("ANS", parent=body_s, leftIndent=12, fontSize=9)
                            ))
                        # Explanation
                        if q.get("explanation"):
                            story.append(Paragraph(
                                f"<i>Explanation: {safe_text(q['explanation'][:200])}</i>",
                                ParagraphStyle("EXP", parent=body_s, leftIndent=12, fontSize=8, textColor=SLATE)
                            ))
                        story.append(Spacer(1, 3*mm))
                else:
                    story.append(Paragraph("No quiz generated yet. Go to the Quiz tab and click Generate.", body_s))

        # ── Study Notes (for when study_notes is selected) ───────────────
        if "study_notes" in ct:
            detailed = [s for s in data["summaries"] if s.summary_type == "detailed"]
            medium   = [s for s in data["summaries"] if s.summary_type == "medium"]
            chosen   = detailed or medium
            if chosen:
                section_header("Study Notes")
                for line in chosen[0].content.split("\n"):
                    line = line.strip()
                    if not line: continue
                    if line.startswith("## "):
                        story.append(Paragraph(safe_text(line[3:]), h2_s))
                    elif line.startswith("• "):
                        story.append(Paragraph(f"• {safe_text(line[2:])}", body_s))
                    else:
                        story.append(Paragraph(safe_text(line), body_s))

        doc.build(story)
        return path

    # ── DOCX ─────────────────────────────────────────────────────────────

    @staticmethod
    def _docx(out_dir: str, project_id: str, data: dict) -> str:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        project = data["project"]

        # Title
        title = doc.add_heading("VideoLM — AI Video Analysis", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        src = project.source_url or project.source_filename or "Video"
        doc.add_paragraph(f"Source: {src}")
        doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        doc.add_paragraph()

        # Summary
        summaries = [s for s in data["summaries"] if s.summary_type == "medium"]
        if summaries:
            doc.add_heading("AI Summary", 1)
            doc.add_paragraph(summaries[0].content)

        # Bullet summary
        bullets = [s for s in data["summaries"] if s.summary_type == "bullets"]
        if bullets:
            doc.add_heading("Key Points", 1)
            for line in bullets[0].content.splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip().lstrip("•").strip(), style="List Bullet")

        # Transcript
        if "transcript" in data["content_types"] and data["transcript"]:
            doc.add_page_break()
            doc.add_heading("Full Transcript", 1)
            t = data["transcript"]
            for seg in (t.segments or [])[:200]:
                ts = seg.get("start", 0)
                m, s = int(ts // 60), int(ts % 60)
                p = doc.add_paragraph()
                run = p.add_run(f"[{m:02d}:{s:02d}] ")
                run.bold = True
                run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
                p.add_run(seg.get("text", ""))

        path = os.path.join(out_dir, f"report_{project_id[:8]}.docx")
        doc.save(path)
        return path

    # ── TXT ──────────────────────────────────────────────────────────────

    @staticmethod
    def _txt(out_dir: str, project_id: str, data: dict) -> str:
        path = os.path.join(out_dir, f"transcript_{project_id[:8]}.txt")
        t = data["transcript"]
        with open(path, "w", encoding="utf-8") as f:
            f.write("VideoLM — Transcript\n")
            f.write("=" * 50 + "\n\n")
            if t:
                for seg in (t.segments or []):
                    ts = seg.get("start", 0)
                    m, s = int(ts // 60), int(ts % 60)
                    f.write(f"[{m:02d}:{s:02d}] {seg.get('text', '').strip()}\n")
        return path

    # ── JSON ─────────────────────────────────────────────────────────────

    @staticmethod
    def _json(out_dir: str, project_id: str, data: dict) -> str:
        path = os.path.join(out_dir, f"export_{project_id[:8]}.json")
        project = data["project"]
        t = data["transcript"]
        output = {
            "generated_at": datetime.utcnow().isoformat(),
            "project": {
                "id": project.id, "title": project.title,
                "source": project.source_url or project.source_filename,
                "duration": project.duration_seconds,
                "language": project.language,
            },
            "transcript": {
                "text": t.full_text if t else "",
                "language": t.language if t else "en",
                "segments": t.segments if t else [],
            } if t else None,
            "summaries": [
                {"type": s.summary_type, "model": s.model_used, "content": s.content}
                for s in data["summaries"]
            ],
            "frames": [
                {
                    "timestamp": f.timestamp_label,
                    "caption": f.caption,
                    "ocr_text": f.ocr_text,
                }
                for f in data["frames"]
            ],
        }
        with open(path, "w", encoding="utf-8") as f:
            json.dump(output, f, indent=2, ensure_ascii=False)
        return path

    # ── SRT ──────────────────────────────────────────────────────────────

    @staticmethod
    def _srt(out_dir: str, project_id: str, data: dict) -> str:
        path = os.path.join(out_dir, f"subtitles_{project_id[:8]}.srt")
        t = data["transcript"]
        with open(path, "w", encoding="utf-8") as f:
            for i, seg in enumerate((t.segments if t else []), 1):
                start = _srt_time(seg.get("start", 0))
                end = _srt_time(seg.get("end", seg.get("start", 0) + 3))
                f.write(f"{i}\n{start} --> {end}\n{seg.get('text', '').strip()}\n\n")
        return path

    # ── VTT ──────────────────────────────────────────────────────────────

    @staticmethod
    def _vtt(out_dir: str, project_id: str, data: dict) -> str:
        path = os.path.join(out_dir, f"subtitles_{project_id[:8]}.vtt")
        t = data["transcript"]
        with open(path, "w", encoding="utf-8") as f:
            f.write("WEBVTT\n\n")
            for i, seg in enumerate((t.segments if t else []), 1):
                start = _vtt_time(seg.get("start", 0))
                end = _vtt_time(seg.get("end", seg.get("start", 0) + 3))
                f.write(f"{i}\n{start} --> {end}\n{seg.get('text', '').strip()}\n\n")
        return path

    # ── CSV ──────────────────────────────────────────────────────────────

    @staticmethod
    def _csv(out_dir: str, project_id: str, data: dict) -> str:
        path = os.path.join(out_dir, f"transcript_{project_id[:8]}.csv")
        t = data["transcript"]
        with open(path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["index", "start", "end", "text", "speaker"])
            for i, seg in enumerate((t.segments if t else []), 1):
                writer.writerow([
                    i, seg.get("start", ""), seg.get("end", ""),
                    seg.get("text", ""), seg.get("speaker", ""),
                ])
        return path

    # ── ZIP bundle ───────────────────────────────────────────────────────

    @staticmethod
    def _zip(out_dir: str, project_id: str, data: dict) -> str:
        # Generate all sub-formats first
        files = []
        for fn, fmt in [
            (ExportService._pdf, "pdf"),
            (ExportService._docx, "docx"),
            (ExportService._txt, "txt"),
            (ExportService._json, "json"),
            (ExportService._srt, "srt"),
        ]:
            try:
                p = fn(out_dir, project_id, data)
                if p and os.path.exists(p):
                    files.append(p)
            except Exception as e:
                logger.warning(f"Sub-export {fmt} failed: {e}")

        zip_path = os.path.join(out_dir, f"videolm_bundle_{project_id[:8]}.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for fp in files:
                zf.write(fp, os.path.basename(fp))
        return zip_path


def _srt_time(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _vtt_time(seconds: float) -> str:
    return _srt_time(seconds).replace(",", ".")